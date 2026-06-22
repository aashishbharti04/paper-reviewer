"""PDF and DOCX text extraction.

PDFs: tries pymupdf (fast) first, falls back to pdfplumber.
DOCX: python-docx.
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Optional  # noqa: F401  (used in type hints below)


MAX_CHARS = 12000  # cap to keep prompt within free-tier model context


def extract_paper_id(filename: str) -> Optional[str]:
    """Pull the paper ID from the filename. Looks for the longest run of digits."""
    stem = Path(filename).stem
    matches = re.findall(r"\d+", stem)
    if not matches:
        return None
    return max(matches, key=len)


def extract_text(path: str) -> str:
    """Extract text from a PDF or DOCX. Returns plain text, truncated to MAX_CHARS."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(path)
    elif ext == ".docx":
        text = _extract_docx(path)
    elif ext == ".doc":
        raise ValueError(
            "Old .doc format not supported. Save the file as .docx or .pdf first."
        )
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    text = _clean(text)
    return text[:MAX_CHARS]


def count_pages(path: str) -> int:
    """Get the page count.

    PDF: exact (via pymupdf).
    DOCX: estimated as word_count / 380 (typical academic page ~350-400 words).
    Returns 0 if it can't be determined.
    """
    ext = Path(path).suffix.lower()
    try:
        if ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(path)
                n = doc.page_count
                doc.close()
                return n
            except ImportError:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    return len(pdf.pages)
        if ext == ".docx":
            import docx
            d = docx.Document(path)
            words = 0
            for p in d.paragraphs:
                words += len(p.text.split())
            for t in d.tables:
                for row in t.rows:
                    for cell in row.cells:
                        words += len(cell.text.split())
            return max(1, round(words / 380))
    except Exception:
        return 0
    return 0


# Headings we look for to assess paper completeness.
# Map of canonical section -> regex patterns. A section counts as present if any of
# its keywords appears ANYWHERE inside a heading line (not just right after the number),
# so "3. System Model and Methodology" correctly counts as a methodology section.
_SECTION_PATTERNS = {
    "abstract":      r"\babstract\b",
    "keywords":      r"\b(?:keywords?|key words?|index terms?)\b",
    "introduction":  r"\bintroduction\b",
    "related_work":  r"\b(?:related works?|literature (?:review|survey)|prior works?|background(?: and related work)?|state of the art)\b",
    "methodology":   r"\b(?:methodolog(?:y|ies)|method[s]?|proposed (?:method|methodology|system|approach|model|framework|scheme|algorithm|technique|work|architecture)|system (?:model|architecture|design|overview|framework)|problem (?:formulation|definition|statement)|mathematical model(?:l?ing)?|materials and methods|experimental (?:setup|design|methodology)|research (?:method|methodology|design)|model(?:l?ing)?|framework|approach|implementation|design and implementation)\b",
    "results":       r"\b(?:results?|experiments?|experimental (?:results|setup|evaluation)|evaluation|performance (?:analysis|evaluation)|findings|results and (?:discussion|analysis)|analysis and results)\b",
    "discussion":    r"\bdiscussions?\b",
    "conclusion":    r"\b(?:conclusions?|concluding remarks|conclusion(?:s)? and future (?:work|scope)|summary and conclusion)\b",
    "references":    r"\b(?:references|bibliography|works cited)\b",
    "acknowledgements": r"\b(?:acknowledg(?:e?ments?|ments?))\b",
}

# A line is treated as a HEADING if it starts with a section number (1, 1.1, IV, A.)
# and is short, OR is a short Title-Case / ALL-CAPS line.
_HEADING_NUM = re.compile(r"^\s*(?:[IVXLC]{1,5}\.?|[A-Z]\.|\d+(?:\.\d+)*\.?)\s+\S")


def _heading_lines(text: str) -> list[str]:
    out = []
    for raw in text.split("\n"):
        ln = raw.strip()
        if not ln or len(ln) > 90:
            continue
        if _HEADING_NUM.match(ln):
            out.append(ln)
        elif len(ln) <= 60 and ln == ln.upper() and any(c.isalpha() for c in ln):
            out.append(ln)  # ALL-CAPS short heading
    return out


def detect_sections(text: str, *, path: Optional[str] = None) -> dict[str, bool]:
    """Boolean flags for each canonical section.

    Strategy: extract heading-like lines, then check whether each section's keyword
    appears anywhere within ANY heading. Falls back to a whole-text line-start scan.
    If `path` is given, re-extracts the FULL document (no MAX_CHARS truncation) so
    end sections (references, acknowledgements) are seen.
    """
    haystack = text
    if path:
        try:
            full = _extract_full_text(path)
            if full and len(full) > len(haystack):
                haystack = full
        except Exception:
            pass

    headings = _heading_lines(haystack)
    headings_blob = "\n".join(headings)

    flags: dict[str, bool] = {}
    for section, pat in _SECTION_PATTERNS.items():
        rx = re.compile(pat, re.IGNORECASE)
        # 1. keyword anywhere inside a heading line (primary, accurate)
        found = bool(rx.search(headings_blob))
        # 2. fallback: keyword right after numbering at the start of any line
        if not found:
            found = bool(re.compile(rf"(?im)^[\s\d\.IVXivx]*{pat}\b").search(haystack))
        flags[section] = found
    return flags


def _extract_full_text(path: str) -> str:
    """Full text extraction without MAX_CHARS truncation. Used only by detect_sections."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(path)
            parts = [page.get_text() for page in doc]
            doc.close()
            return _clean("\n".join(parts))
        except ImportError:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                return _clean("\n".join((p.extract_text() or "") for p in pdf.pages))
    if ext == ".docx":
        return _extract_docx(path)
    return ""


def _extract_pdf(path: str) -> str:
    try:
        import fitz  # pymupdf

        doc = fitz.open(path)
        parts = []
        for page in doc:
            parts.append(page.get_text())
            if sum(len(p) for p in parts) > MAX_CHARS * 2:
                break
        doc.close()
        return "\n".join(parts)
    except ImportError:
        pass
    try:
        import pdfplumber

        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                parts.append(t)
                if sum(len(p) for p in parts) > MAX_CHARS * 2:
                    break
        return "\n".join(parts)
    except ImportError as e:
        raise RuntimeError(
            "Neither pymupdf nor pdfplumber installed. Run install.bat."
        ) from e


def _extract_docx(path: str) -> str:
    try:
        import docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx not installed. Run install.bat."
        ) from e

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _clean(text: str) -> str:
    """Collapse excessive whitespace and remove non-printable junk."""
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[\t\f\v]+", " ", text)
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def guess_title(text: str) -> str:
    """Best-effort title extraction: longest non-trivial line from the first 1500 chars."""
    head = text[:1500]
    lines = [ln.strip() for ln in head.split("\n") if ln.strip()]
    candidates = [ln for ln in lines if 20 <= len(ln) <= 200 and not ln.lower().startswith(("abstract", "keywords", "introduction"))]
    if not candidates:
        return ""
    return max(candidates[:10], key=len)
